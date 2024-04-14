from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import requests
from bs4 import BeautifulSoup
import random
import undetected_chromedriver as uc
from seleniumwire import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException
import os
from docx import *
from fake_useragent import UserAgent
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

class ClubExtract():
    def __init__(self):
        self.driver = ''
        self.list_of_urls = []
        self.project_club_url = "https://www.ualberta.ca/engineering/student-life/clubs-projects-competitions.html"
    
    def getProxies(self):
        r = requests.get('https://free-proxy-list.net/', verify=False)
        soup = BeautifulSoup(r.content, 'html.parser')
        table = soup.find('tbody')
        total_rows = len(table)
        with open("proxies.txt", 'w') as file:
            for index, row in enumerate(table):
                if row.find_all('td')[4].text == 'elite proxy' and row.find_all('td')[5].text == 'yes':
                    proxy = ':'.join([row.find_all('td')[0].text, row.find_all('td')[1].text])
                    file.write(proxy)
                    if index < total_rows - 1:
                        file.write('\n')

    def setupDriver(self):

        proxies = []
        options = Options()

        user_agent = UserAgent()

        user_string = user_agent.random

        with open("proxies.txt", 'r') as file:
            for line in file:
                proxies.append(line)
                options.add_argument(f'--proxy-server={line}')
        options.add_argument('--headless')

        random_proxy = random.choice(proxies)      

        seleniumwire_options = {
            'proxy': {
                'http': f'{random_proxy}',
                'https': f'{random_proxy}',
                'verify_ssl': False
            },
        }

        chrome_options = webdriver.ChromeOptions()
        chrome_options.add_argument(f'--user-agent={user_string}')
        chrome_options.add_argument('--disable-blink-features=AutomationControlled')
        chrome_options.headless = True
        self.driver = uc.Chrome(options=chrome_options, seleniumwire_options=seleniumwire_options)
    
    def extract_clubs(self):
        self.driver.get("https://alberta.campuslabs.ca/engage/organizations")
        org_results = self.driver.find_element(By.ID, "org-search-results")

        for i in range(0, 48):
            outlined_button = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable(self.driver.find_element(By.CLASS_NAME, "outlinedButton")))
            outlined_button.click()
        
        orgs_a_tags = org_results.find_elements(By.TAG_NAME, 'a')

        self.list_of_urls = [a_tag.get_attribute("href") for a_tag in orgs_a_tags]
        
        print(self.list_of_urls)

    
    def write_pdf(self):
        document = Document()

        document.add_heading('University of Alberta Clubs', 0)

        for url in self.list_of_urls:
            self.setupDriver()
            self.driver.get(url)

            paragraph_div = ''
            club_name = ''
            try:
                club_name = self.driver.find_element(By.TAG_NAME, "h1").text
                paragraph_div = self.driver.find_element(By.CLASS_NAME, "bodyText-large.userSupplied")
                document.add_heading(club_name, level=1)
                description_paragraph = document.add_paragraph(style='List Bullet')
                run = description_paragraph.add_run("Club Description:")
                run.bold = True

                all_pars = paragraph_div.find_elements(By.TAG_NAME, "p")

            except NoSuchElementException:
                paragraph = "No club information available"
                all_pars = []

            if all_pars:
                for i in range(0, len(all_pars)):
                    if i == 0:
                        description_paragraph.add_run('\n' + all_pars[i].text + "\n")
                    else:
                        description_paragraph.add_run(all_pars[i].text + "\n")
            else:
                description_paragraph.add_run('\n' + "No description found" + '\n')
        

            os.system("taskkill /im chrome.exe /f")

            # Engineering project club extraction
        
        self.setupDriver()
        self.driver.get(self.project_club_url)

        group_container = self.driver.find_element(By.ID, "accordionList112")

        all_project_clubs = group_container.find_elements(By.CLASS_NAME, "card")

        for i in range(0, len(all_project_clubs)):

            html_content = all_project_clubs[i].get_attribute("outerHTML")

            # Parse the HTML content with BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")

            # Find the div with class "card-body"
            card_body_div = soup.find("div", class_="card-body")

            
            project_club = card_body_div.find("h2").text

            
            document.add_heading(project_club, level=1)

            description_paragraph = document.add_paragraph(style='List Bullet')
            run = description_paragraph.add_run("Club Description:")
            run.bold = True

            project_club_paragraphs = card_body_div.find_all("p")

            if project_club_paragraphs:
                for i in range(0, len(project_club_paragraphs)):
                    if i == 0:
                        description_paragraph.add_run("\n" + project_club_paragraphs[i].text + "\n")
                        print(description_paragraph)
                    else:
                        description_paragraph.add_run(project_club_paragraphs[i].text + "\n")
                        print(description_paragraph)
            else:
                description_paragraph.add_run('\n' + "No description found" + '\n')
            
        os.system("taskkill /im chrome.exe /f")

        document.save("list_of_clubs.docx")
        self.driver.quit()

    def run(self):
        self.getProxies()
        self.setupDriver()
        self.extract_clubs()
        self.write_pdf()

if __name__ == "__main__":
    extract_object = ClubExtract() # can put compe, software, or nano i    n the constructor
    # extract_object.course_description_extract('ece 321')

    extract_object.run()