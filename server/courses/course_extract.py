from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import requests
from bs4 import BeautifulSoup
import random
import undetected_chromedriver as uc
from seleniumwire import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException
import os
from docx import *
from googlesearch import search
from fake_useragent import UserAgent
from reddit_extract import RedditExtract
from docx.enum.text import WD_BREAK

class CourseExtract():

    def __init__(self):
        self.engineering_url = ['https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55192&returnto=13670']
        self.list_of_file_paths = ['courses_compe\compe_group2_electives.txt']
        self.driver = ''
    
    def getProxies(self):
        r = requests.get('https://free-proxy-list.net/', verify=False)
        soup = BeautifulSoup(r.content, 'html.parser')
        table = soup.find('tbody')
        total_rows = len(table)
        with open("proxies.txt", 'w') as file:
            for index, row in enumerate(table):
                print(row.find_all('td')[4].text)
                print(row.find_all('td')[5].text == 'yes')
                # row.find_all('td')[4].text == 'elite proxy' and 
                if row.find_all('td')[5].text == 'yes':
                    proxy = ':'.join([row.find_all('td')[0].text, row.find_all('td')[1].text])
                    file.write(proxy)
                    if index < total_rows - 1:
                        file.write('\n')

    def setupDriver(self):

        proxies = []
        options = Options()

        user_agent = UserAgent()

        user_string = user_agent.random

        with open("proxies.txt", 'r') as file:
            for line in file:
                proxies.append(line)
                options.add_argument(f'--proxy-server={line}')
        options.add_argument('--headless')

        random_proxy = random.choice(proxies)      

        seleniumwire_options = {
            'proxy': {
                'http': f'{random_proxy}',
                'https': f'{random_proxy}',
                'verify_ssl': False
            },
        }

        chrome_options = webdriver.ChromeOptions()
        chrome_options.add_argument(f'--user-agent={user_string}')
        chrome_options.add_argument('--disable-blink-features=AutomationControlled')
        chrome_options.headless = True
        self.driver = uc.Chrome(options=chrome_options, seleniumwire_options=seleniumwire_options)
    
    def extract_group_two(self):

        url_to_search = ''
        list_of_non_grp2s = []

        list_of_non_grp2_compe_normal = ['CMPUT 274', 'ECE 202', 'ECE 210', 'ENGG 299', 'MATH 201', 'MATH 209', 'CMPUT 272', 'CMPUT 275', 'ECE 212', 'ECE 240', 'PHYS 230', 'WKEXP 901', 'ECE 311', 'ECE 325', 'STAT 235', 'WKEXP 902', 'WKEXP 903', 'CMPUT 291', 'CMPUT 301', 'CMPUT 379', 'ECE 322', 'ECE 315', 'ECE 487', 'ENGG 404', 'WKEXP 904', 'WKEXP 905', 'ECE 420', 'ECE 422', 'ECE 493', 'ENGG 400','ECE 203', 'ECE 302', 'ECE 340', 'ECE 304', 'ECE 342', 'ECE 410', 'ECE 492']
        list_of_non_grp2_compe_nano = ['CMPUT 274', 'ECE 202', 'ECE 210', 'ENGG 299', 'MATH 201', 'MATH 209', 'CMPUT 272', 'CMPUT 275', 'ECE 203', 'ECE 212', 'ECE 240', 'PHYS 230', 'WKEXP 901', 'ECE 302', 'ECE 311', 'ECE 325', 'WKEXP 902', 'WKEXP 903', 'CMPUT 291', 'ECE 304', 'ECE 342', 'ECE 410', 'ENGG 404', 'CMPUT 301', 'ECE 315', 'ECE 403', 'ECE 450', 'ECE 475', 'WKEXP 904', 'WKEXP 905', 'ECE 412', 'ECE 457', 'ECE 492', 'ENGG 400']
        list_of_non_grp2_software = ['CMPUT 274', 'ECE 202', 'ECE 210', 'ENGG 299', 'MATH 201', 'MATH 209', 'CMPUT 272', 'CMPUT 275', 'ECE 212', 'ECE 240', 'PHYS 230', 'WKEXP 901', 'ECE 311', 'ECE 321', 'ECE 325', 'STAT 235', 'WKEXP 902', 'WKEXP 903', 'CMPUT 291', 'CMPUT 301', 'CMPUT 379', 'ECE 322', 'ECE 315', 'ECE 421', 'ECE 487', 'ENGG 404', 'WKEXP 904', 'WKEXP 905', 'ECE 420', 'ECE 422', 'ECE 493', 'ENGG 400']
        list_of_non_grp2_ee_bio = ['ECE 202', 'ECE 210', 'ENGG 299' 'ECE 210', 'MATH 201', 'MATH 209', 'ECE 203', 'ECE 212', 'ECE 220', 'ECE 240', 'PHYS 230', 'BIOL 107', 'ECE 302', 'ECE 312', 'ECE 340', 'ECE 370', 'MATH 309', 'ECE 303', 'ECE 342', 'ECE 360', 'ECE 380', 'ENGG 404', 'ECE 405', 'ECE 440', 'ECE 490', 'PHYSL 210', 'ECE 491', 'ENGG 400', 'PHYSL 210']
        list_of_non_grp2_ee_normal = ['ECE 202', 'ECE 210', 'ENGG 299', 'MATH 201', 'MATH 209', 'ECE 203', 'ECE 212', 'ECE 220', 'ECE 240', 'PHYS 230', 'WKEXP 901', 'ECE 302', 'ECE 312', 'ECE 330', 'ECE 340', 'ECE 370', 'MATH 309', 'WKEXP 902', 'WKEXP 903', 'ECE 303', 'ECE 332', 'ECE 342', 'ECE 360', 'ECE 380', 'WKEXP 904', 'WKEXP 905', 'ECE 490', 'ENGG 404', 'ECE 491', 'ENGG 400']
        list_of_non_grp2_ee_nano = ['ECE 202', 'ECE 210', 'ENGG 299', 'MATH 201', 'MATH 209', 'ECE 203', 'ECE 212', 'ECE 220', 'ECE 240', 'PHYS 230', 'WKEXP 901', 'ECE 302', 'ECE 312', 'ECE 340', 'ECE 370', 'ECE 471', 'MATH 309', 'WKEXP 902', 'WKEXP 903', 'WKEXP 904', 'ECE 303', 'ECE 341', 'ECE 342', 'ECE 456', 'WKEXP 905', 'ECE 360', 'ECE 457', 'ECE 490', 'ENGG 404', 'ECE 450', 'ECE 475', 'ECE 491', 'ENGG 400']
        list_of_non_grp2_enphys_normal = ['ECE 202', 'ENGG 299', 'MATH 201', 'MATH 209', 'PHYS 281', 'PHYS 292', 'ECE 203', 'ECE 220', 'ECE 240', 'PHYS 271', 'PHYS 292', 'WKEXP 901', 'ECE 210', 'ECE 302', 'ECE 340', 'ECE 471', 'MATH 311', 'PHYS 381', 'WKEXP 902', 'WKEXP 903', 'WKEXP 904', 'ECE 303', 'ECE 341', 'PHYS 244', 'PHYS 311', 'PHYS 372', 'WKEXP 905', 'ECE 494', 'ENGG 404', 'PHYS 415', 'PHYS 481', 'ECE 360', 'ECE 495', 'ENGG 400']
        list_of_non_grp2_enphys_nano = ['ECE 202', 'ENGG 299', 'MATH 201', 'MATH 209', 'PHYS 281', 'PHYS 292', 'ECE 203', 'ECE 240', 'PHYS 244', 'PHYS 271', 'PHYS 292', 'WKEXP 901', 'ECE 210', 'ECE 302', 'ECE 457', 'ECE 471', 'MATH 311', 'PHYS 381', 'WKEXP 902', 'WKEXP 903', 'WKEXP 904', 'ECE 303', 'ECE 341', 'ECE 360', 'ECE 456', 'PHYS 311', 'PHYS 372', 'WKEXP 905', 'CHEM 261', 'ECE 494', 'ENGG 404', 'PHYS 415', 'PHYS 481', 'BIOCH 200', 'ECE 455', 'ECE 495', 'ENGG 400']
        
        list_of_all_grp2s = []
        for i in range(0, len(self.engineering_url)):

            self.driver.get(self.engineering_url[i])

            h1 = self.driver.find_elements(By.CLASS_NAME, 'acalog-course')
            list_of_grp2 = []

            if self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55199&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_software
            elif self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55192&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_compe_normal
            elif self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55194&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_compe_nano
            elif self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55167&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_ee_bio
            elif self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55178&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_ee_normal
            elif self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55180&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_ee_nano
            elif self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55200&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_enphys_normal
            elif self.engineering_url[i] == "https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55203&returnto=13670":
                list_of_non_grp2s = list_of_non_grp2_enphys_nano
            
            
            
            for i in range(0, len(h1)):
                split_elective = h1[i].text.split()
                try:
                    value = int(split_elective[1])
                    element = (split_elective[0] + ' ' + split_elective[1])
                    if (element.strip()) not in list_of_non_grp2s:
                        if i<len(h1)-1:
                            list_of_grp2.append(element.strip() + "\n")
                        else:
                            list_of_grp2.append(element.strip())

                except ValueError:
                    pass
            
            list_of_all_grp2s.append(list_of_grp2)
            list_of_grp2 = []

        self.driver.quit()
        os.system("taskkill /im chrome.exe /f")
        return list_of_all_grp2s
    
    def create_grp2_text(self):

        list_of_all_grp2s = self.extract_group_two()

        file_path = ''

        for i in range(len(list_of_all_grp2s)):
            file_path = self.list_of_file_paths[i]
            if not os.path.exists(file_path):
                with open(file_path, 'w') as file:
                    last_element = list_of_all_grp2s[i][-1]  # Get the last element
                    if last_element.endswith('\n'):  # Check if it ends with '\n'
                        last_element = last_element.rstrip('\n')  # Remove trailing newline character if exists
                        list_of_all_grp2s[i][-1] = last_element  # Set the modified last element
                    file.writelines(list_of_all_grp2s[i])

    def course_description_extract(self, course_name):

        split_course = course_name.split()
        course_url = 'https://apps.ualberta.ca/catalogue/course/' + split_course[0].lower() + '/' + split_course[1]

        self.driver.get(course_url)

        course_info = (self.driver.find_element("xpath", "//div[@class='container']/p[2]").text).split('.')

        term_and_instructor_info = (self.driver.find_elements(By.CLASS_NAME, 'mb-5'))

        prerequisites = ""

        for i in range(0, len(course_info)):
            if "Prerequisite" in course_info[i]:
                prerequisites = course_info.pop(i)
                break
        
        course_description = '.'.join(course_info) # All course info sans prereqs

        
        term_and_instructor_info = self.driver.find_elements(By.CLASS_NAME, 'mb-5')

        term_and_profs = {}

        for element in term_and_instructor_info:

            what_term = element.find_element(By.TAG_NAME, 'h2').text
            
            try:
                table_elements = element.find_element(By.TAG_NAME, 'table')
                td_elements = table_elements.find_elements(By.TAG_NAME, 'td')
                
                # Process the table data
            except NoSuchElementException:
                term_and_profs[what_term] = []
                continue
            
            for td_element in td_elements:
                try:
                    data_card_title = td_element.get_attribute('data-card-title') # Check if the td element has the data-card-title attribute
                    if data_card_title == "Instructor(s)":

                        a_tags = td_element.find_elements(By.TAG_NAME, 'a')
                        if what_term not in term_and_profs:
                            term_and_profs[what_term] = []

                        for atag in a_tags:
                            term_and_profs[what_term].append(atag.text)
                except:
                    pass
        
        os.system("taskkill /im chrome.exe /f")
        self.driver.quit()
        return term_and_profs, prerequisites, course_description
    
    def extract_prof(self, professor):
        self.setupDriver()
        first_and_last_names = professor.split()

        professor_to_find = f"https://www.google.com/search?q={'+'.join(first_and_last_names)}+University+of+Alberta+Rate+My+Professor"
        self.driver.get(professor_to_find)
        professor_url = ""

        try:
            regular_results = self.driver.find_elements(By.XPATH, "//div[@class='g']")

            for result in regular_results:
                link_element = result.find_element(By.TAG_NAME, 'a')
                professor_url = link_element.get_attribute("href")
                professor_url, _, _ = professor_url.partition('#')  # professor_url after partitioning. _, only 
                break
        except NoSuchElementException:
            print("No regular search results found")

        if not professor_url:
            try:
                alt_results = self.driver.find_elements(By.CLASS_NAME, "MjjYud")

                # Iterate over alternative results to find the link
                for result in alt_results:
                    link_element = result.find_element(By.TAG_NAME, 'a')
                    professor_url = link_element.get_attribute("href")
                    break  # Stop after getting the first link
            except NoSuchElementException:
                print("No alternative search results found")
        # Need to verify if the url is valid (the proper associated rate my prof rating for the professor in question)
                
        ratemyprof_rating = ""
        
        if "ratemyprofessors" not in professor_url:
            ratemyprof_rating = "The professor does not have a rating on Rate My Professor"
        else:
            self.driver.get(professor_url)
            try:
                first_name_rmp_div = self.driver.find_element(By.CLASS_NAME, "NameTitle__Name-dowf0z-0.cfjPUG")
                first_name_rmp = first_name_rmp_div.find_element(By.TAG_NAME, "span").text
                last_name_rmp = self.driver.find_element(By.CLASS_NAME, "NameTitle__LastNameWrapper-dowf0z-2.glXOHH").text

                print(first_name_rmp, last_name_rmp)

                if first_name_rmp.lower()!=first_and_last_names[0].lower() or last_name_rmp.lower()!=first_and_last_names[1].lower():
                    ratemyprof_rating = "The professor does not have a rating on Rate My Professor"
                else:
                    prof_rating = self.driver.find_element(By.CLASS_NAME, "RatingValue__Numerator-qw8sqy-2.liyUjw").text

                    print(prof_rating)
                    ratemyprof_rating = prof_rating + "/5"
            
            except NoSuchElementException:
                ratemyprof_rating = "The professor does not have a rating on Rate My Professor"

                print(ratemyprof_rating)
        
        os.system("taskkill /im chrome.exe /f")
        self.driver.quit()
            
        return ratemyprof_rating


    def write_pdf(self):

        for discipline in self.list_of_file_paths:
            document = Document()

            if self.engineering_url[0] == 'https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55203&returnto=13670' or self.engineering_url[0] == 'https://calendar.ualberta.ca/preview_program.php?catoid=44&poid=55200&returnto=13670':
                document.add_heading('Program & Technical Electives', 0)
            else:
                document.add_heading('Group 2 Electives', 0)

            with open(discipline, "r") as file:

                for course in file:
                    if course.strip() == "":
                        exit()
                    self.setupDriver()
                    terms_and_profs, prerequisites, course_description = self.course_description_extract(course)
                    document.add_heading(course, level=1)
                    description_paragraph = document.add_paragraph(style='List Bullet')
                    run = description_paragraph.add_run("Course Description:")
                    run.bold = True
                    
                    # Add the course description to the parent paragraph
                    description_paragraph.add_run('\n' + course_description + "\n")

                    # Adding prereqs
                    print(terms_and_profs)
                    prerequisites_paragraph = document.add_paragraph(style='List Bullet')
                    if "Prerequisites" not in prerequisites and len(prerequisites.strip()) > 0:
                        run_prerequisites = prerequisites_paragraph.add_run("Prerequisite:" + "\n")
                        run_prerequisites.bold = True
                        prerequisites_paragraph.add_run(prerequisites.replace("Prerequisite: ", "") + "\n")
                    elif "Prerequisites" in prerequisites and len(prerequisites.strip()) > 0:
                        run_prerequisites = prerequisites_paragraph.add_run("Prerequisites:" + "\n")
                        run_prerequisites.bold = True
                        prerequisites_paragraph.add_run(prerequisites.replace("Prerequisites: ", "") + "\n")
                    elif len(prerequisites.strip()) == 0:
                        run_prerequisites = prerequisites_paragraph.add_run("Prerequisites:" + "\n")
                        run_prerequisites.bold = True
                        prerequisites_paragraph.add_run("None" + "\n")
                    
                    # Adding terms
                    terms_paragraph = document.add_paragraph(style='List Bullet')
                    run_terms = terms_paragraph.add_run("Terms the course is available in:" + "\n")
                    run_terms.bold = True

                    keys = list(terms_and_profs.keys())

                    if len(keys)>0:
                        for i, key in enumerate(keys):
                            terms_paragraph.add_run(key)
                            if i < len(keys) - 1:
                                terms_paragraph.add_run(", ")
                            else:
                                terms_paragraph.add_run("\n")
                    else:
                        terms_paragraph.add_run("No term decided yet/not offered this year" + "\n")
                    
                    # Adding professors
                    professors_paragraph = document.add_paragraph(style='List Bullet')
                    run_terms = professors_paragraph.add_run("Instructor(s):" + "\n")
                    run_terms.bold = True

                    counter = 0

                    if len(terms_and_profs) > 0:
                        for term, instructors in terms_and_profs.items():
                            counter += 1
                            if counter < len(terms_and_profs):
                                if len(instructors) == 1:
                                    professors_paragraph.add_run(instructors[0] + " (teaching in " + term + "), ")
                                elif len(instructors) > 1:

                                    # Convert list to set and then back again
                                    unique_instructors = list(set(instructors))
                                    for i in range(len(unique_instructors)):
                                        if i == len(unique_instructors) - 1:
                                            professors_paragraph.add_run(unique_instructors[i] + " (teaching in " + term + "), ")
                                            break
                                        else:
                                            professors_paragraph.add_run(unique_instructors[i] + " (teaching in " + term + "), ")
                                else:
                                    professors_paragraph.add_run("Instructor(s) undecided for " + term + ", ")
                            else:
                                if len(instructors) == 1:
                                    professors_paragraph.add_run(instructors[0] + " (teaching in " + term + "), ")
                                elif len(instructors) > 1:
                                    # Remove duplicates using set and convert back to list
                                    unique_instructors = list(set(instructors))
                                    for i in range(len(unique_instructors)):
                                        if i == len(unique_instructors) - 1:
                                            professors_paragraph.add_run(unique_instructors[i] + " (teaching in " + term + ")")
                                            break
                                        else:
                                            professors_paragraph.add_run(unique_instructors[i] + " (teaching in " + term + "), ")
                                else:
                                    professors_paragraph.add_run("Instructor(s) undecided for " + term + "")
                    else:
                        professors_paragraph.add_run("No instructor teaching the course")
                    
                    professors_paragraph.add_run("\n")

                    
                    # Adding professor rating
                        
                    rating_paragraph = document.add_paragraph(style='List Bullet')
                    run_terms = rating_paragraph.add_run("Instructor ratings:" + "\n")
                    run_terms.bold = True

                    list_of_profs = list(terms_and_profs.values())
                    unique_profs = []

                    for sublist in list_of_profs:
                        unique_sublist = list(set(sublist))

                        if len(unique_sublist) > 0:
                            unique_profs.extend(unique_sublist)
                    
                    if len(unique_profs) == 0:
                        rating_paragraph.add_run("No professors teaching this term, so no ratings available at all")
                    else:
                        for i in range(len(unique_profs)):
                            professor = unique_profs[i]
                            rating = self.extract_prof(professor)
                            if i<len(unique_profs) - 1:
                                if rating != "The professor does not have a rating on Rate My Professor":
                                    rating_paragraph.add_run(f"{professor}'s Rate My Professor rating is {rating}, ")
                                else:
                                    rating_paragraph.add_run(rating)
                            else:
                                if rating != "The professor does not have a rating on Rate My Professor":
                                    rating_paragraph.add_run(f"{professor}'s Rate My Professor rating is {rating}")
                                else:
                                    rating_paragraph.add_run(rating)
                    
                    rating_paragraph.add_run("\n")
                    os.system("taskkill /im chrome.exe /f")

                    self.setupDriver()

                    course_rating_object = RedditExtract(self.driver)
                
                    course_rating = course_rating_object.llm_opinion(course.strip())

                    course_difficulty_paragraph = document.add_paragraph(style='List Bullet')
                    course_difficulty_run = course_difficulty_paragraph.add_run("Course Difficulty:" + "\n")
                    course_difficulty_run.bold = True

                    course_difficulty_paragraph.add_run(course_rating)
                    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

                    os.system("taskkill /im chrome.exe /f")
                    self.driver.quit()
                    


            document.save(discipline.replace(".txt", ".docx"))
            self.driver.quit()
            

    def run(self):
        self.getProxies()
        self.setupDriver()
        self.create_grp2_text()
        self.write_pdf()
    
    def run_experimental(self):
        self.getProxies()
        self.setupDriver()
        self.extract_group_two()

if __name__ == "__main__":
    extract_object = CourseExtract()

    extract_object.run()